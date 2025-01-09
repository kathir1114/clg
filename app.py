# from flask import Flask, request, jsonify, send_file
# from flask_sqlalchemy import SQLAlchemy
# from flask_cors import CORS
# from fpdf import FPDF
# from docx import Document
# import os
# import io

# app = Flask(__name__)
# CORS(app)

# # Configure the database
# app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///forms.db'
# app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# db = SQLAlchemy(app)

# # Define database models
# class Form1(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     name = db.Column(db.String(100))
#     mobile_number = db.Column(db.String(10))
#     email = db.Column(db.String(100))
#     date = db.Column(db.String(50))
#     dept = db.Column(db.String(50))
#     photo_filename = db.Column(db.String(100))
#     download_type = db.Column(db.String(10))

# class Form2(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     event_name = db.Column(db.String(100))
#     event_date = db.Column(db.String(50))
#     dept = db.Column(db.String(50))
#     download_type = db.Column(db.String(10))

# # Initialize the database
# with app.app_context():
#     db.create_all()

# # Set up paths for uploaded and downloaded files
# BASE_UPLOAD_PATH = os.path.join(os.getcwd(), 'uploads')
# BASE_DOWNLOAD_PATH = os.path.join(os.getcwd(), 'downloads')

# if not os.path.exists(BASE_UPLOAD_PATH):
#     os.makedirs(BASE_UPLOAD_PATH)

# if not os.path.exists(BASE_DOWNLOAD_PATH):
#     os.makedirs(BASE_DOWNLOAD_PATH)

# # Helper functions to generate PDF, Word, Text files (unchanged)

# # Helper function to generate a PDF
# def generate_pdf(data, filename):
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.cell(200, 10, txt="Form Submission Details", ln=True, align='C')
#     pdf.ln(10)

#     for key, value in data.items():
#         if key == "photoName":
#             pdf.cell(200, 10, txt="Photo Details:", ln=True, align='L')
#             pdf.cell(200, 10, txt=f"  Name: {value}", ln=True, align='L')
#             pdf.cell(200, 10, txt=f"  Size: {data.get('photoSize', 'N/A')}", ln=True, align='L')
#             pdf.cell(200, 10, txt=f"  Type: {data.get('photoType', 'N/A')}", ln=True, align='L')
#         elif key not in ["photoSize", "photoType", "downloadType", "formType"]:
#             pdf.cell(200, 10, txt=f"{key.replace('_', ' ').title()}: {value}", ln=True, align='L')

#     pdf_file_path = os.path.join(BASE_DOWNLOAD_PATH, filename)
#     pdf.output(pdf_file_path)
#     return pdf_file_path

# # Helper function to generate a Word file
# def generate_word(data, filename):
#     doc = Document()
#     doc.add_heading("Form Submission Details", level=1)

#     for key, value in data.items():
#         if key == "photoName":
#             doc.add_heading("Photo Details:", level=2)
#             doc.add_paragraph(f"  Name: {value}")
#             doc.add_paragraph(f"  Size: {data.get('photoSize', 'N/A')}")
#             doc.add_paragraph(f"  Type: {data.get('photoType', 'N/A')}")
#         elif key not in ["photoSize", "photoType", "downloadType", "formType"]:
#             doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

#     word_file_path = os.path.join(BASE_DOWNLOAD_PATH, filename)
#     doc.save(word_file_path)
#     return word_file_path

# # Helper function to generate a Text file
# def generate_text(data, filename):
#     text_file_path = os.path.join(BASE_DOWNLOAD_PATH, filename)
#     with open(text_file_path, 'w') as f:
#         f.write("Form Submission Details\n")
#         f.write("=" * 30 + "\n")

#         for key, value in data.items():
#             if key == "photoName":
#                 f.write("Photo Details:\n")
#                 f.write(f"  Name: {value}\n")
#                 f.write(f"  Size: {data.get('photoSize', 'N/A')}\n")
#                 f.write(f"  Type: {data.get('photoType', 'N/A')}\n")
#             elif key not in ["photoSize", "photoType", "downloadType", "formType"]:
#                 f.write(f"{key.replace('_', ' ').title()}: {value}\n")

#     return text_file_path

# @app.route('/submit-form', methods=['POST'])
# def submit_form():
#     # Handle the form submission
#     form_data = request.form.to_dict()  # Get form data as a dictionary
    
#     # Optional: Handle file upload if there's a file (photo)
#     photo = request.files.get('photo')  # Get file from the form
#     if photo:
#         photo_filename = os.path.join(BASE_UPLOAD_PATH, photo.filename)
#         photo.save(photo_filename)
#         form_data['photoName'] = photo.filename  # Add photo filename to form data

#     download_type = form_data.get('downloadType')  # Get download type from form data

#     # Log for debugging
#     print(f"Requested download type: {download_type}")
#     print(f"Form data: {form_data}")

#     if not download_type:
#         return jsonify({"error": "Missing 'downloadType' field in form data"}), 400

#     try:
#         if download_type == 'pdf':
#             filename = 'form_submission.pdf'
#             file_path = generate_pdf(form_data, filename)
#         elif download_type == 'word':
#             filename = 'form_submission.docx'
#             file_path = generate_word(form_data, filename)
#         elif download_type == 'txt':
#             filename = 'form_submission.txt'
#             file_path = generate_text(form_data, filename)
#         else:
#             return jsonify({"error": "Invalid download type"}), 400

#         return send_file(file_path, as_attachment=True)

#     except Exception as e:
#         print(f"Error generating file: {str(e)}")
#         return jsonify({"error": "An error occurred while generating the file"}), 500

# @app.route('/download/<path:filename>')
# def download_file(filename):
#     file_path = os.path.join(BASE_DOWNLOAD_PATH, filename)
#     if os.path.exists(file_path):
#         return send_file(file_path, as_attachment=True)
#     return jsonify({"error": "File not found"}), 404

# if __name__ == '__main__':
#     app.run(debug=True)

from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import os

app = Flask(__name__)
CORS(app)

# Configure the database
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///forms.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# Define database models
class Form1(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100))
    mobile_number = db.Column(db.String(10))
    email = db.Column(db.String(100))
    date = db.Column(db.String(50))
    dept = db.Column(db.String(50))
    photo_filename = db.Column(db.String(100))
    download_type = db.Column(db.String(10))

# Initialize the database
with app.app_context():
    db.create_all()

# Set up paths for uploaded and downloaded files
BASE_UPLOAD_PATH = os.path.join(os.getcwd(), 'uploads')
BASE_DOWNLOAD_PATH = os.path.join(os.getcwd(), 'downloads')

if not os.path.exists(BASE_UPLOAD_PATH):
    os.makedirs(BASE_UPLOAD_PATH)

if not os.path.exists(BASE_DOWNLOAD_PATH):
    os.makedirs(BASE_DOWNLOAD_PATH)

# Helper function to generate a PDF
def generate_pdf(data, photo_path, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt="Form Submission Details", ln=True, align='C')
    pdf.ln(10)

    for key, value in data.items():
        if key != "photo":
            pdf.cell(200, 10, txt=f"{key.replace('_', ' ').title()}: {value}", ln=True, align='L')

    if photo_path:
        pdf.ln(10)
        pdf.cell(200, 10, txt="Uploaded Photo:", ln=True, align='L')
        pdf.image(photo_path, x=10, y=pdf.get_y(), w=100)

    pdf_file_path = os.path.join(BASE_DOWNLOAD_PATH, filename)
    pdf.output(pdf_file_path)
    return pdf_file_path

# Helper function to generate a Word file
def generate_word(data, photo_path, filename):
    doc = Document()
    doc.add_heading("Form Submission Details", level=1)

    for key, value in data.items():
        if key != "photo":
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    if photo_path:
        doc.add_heading("Uploaded Photo:", level=2)
        doc.add_picture(photo_path, width=Inches(2))

    word_file_path = os.path.join(BASE_DOWNLOAD_PATH, filename)
    doc.save(word_file_path)
    return word_file_path

# Helper function to generate a Text file
def generate_text(data, photo_path, filename):
    text_file_path = os.path.join(BASE_DOWNLOAD_PATH, filename)
    with open(text_file_path, 'w') as f:
        f.write("Form Submission Details\n")
        f.write("=" * 30 + "\n")

        for key, value in data.items():
            if key != "photo":
                f.write(f"{key.replace('_', ' ').title()}: {value}\n")

        if photo_path:
            f.write("\nUploaded Photo: See attached image\n")

    return text_file_path

@app.route('/submit-form', methods=['POST'])
def submit_form():
    # Retrieve form data
    form_data = {
        "name": request.form.get('name'),
        "mobile_number": request.form.get('mobile_number'),
        "email": request.form.get('email'),
        "date": request.form.get('date'),
        "dept": request.form.get('dept')
    }

    download_type = request.form.get('downloadType', 'pdf')
    photo = request.files.get('photo')

    if not photo:
        return jsonify({"error": "No photo uploaded"}), 400

    # Save the photo to the upload directory
    photo_filename = photo.filename
    photo_path = os.path.join(BASE_UPLOAD_PATH, photo_filename)
    photo.save(photo_path)

    # Generate the requested file type
    filename = f"form_submission.{download_type}"
    if download_type == 'pdf':
        file_path = generate_pdf(form_data, photo_path, filename)
    elif download_type == 'word':
        file_path = generate_word(form_data, photo_path, filename)
    elif download_type == 'txt':
        file_path = generate_text(form_data, photo_path, filename)
    else:
        return jsonify({"error": "Unsupported download type"}), 400

    # Send the file to the user
    return send_file(file_path, as_attachment=True)

@app.route('/download/<path:filename>')
def download_file(filename):
    file_path = os.path.join(BASE_DOWNLOAD_PATH, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({"error": "File not found"}), 404

if __name__ == '__main__':
    app.run(debug=True)
